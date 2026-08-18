"""
验证表格行距修复：确保 `w:spacing` 在 `w:ind` 之前（符合 OOXML 规范顺序）。
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from backend.formatter.table import _apply_cell_line_spacing, _set_cell_alignment
from backend.formatter.data_model import TableConfig


def _build_config(**overrides) -> TableConfig:
    """Helper: build a TableConfig with minimal fields."""
    defaults = dict(
        line_spacing=1.5,
        line_spacing_mode="multiple",
        line_spacing_unit="pt",
        cell_align_h="left",
        cell_align_v="center",
        indent_type="none",
        indent_value=0.0,
        indent_unit="字符",
    )
    defaults.update(overrides)
    return TableConfig(**defaults)


def _create_table_doc() -> Document:
    """Create a simple doc with one table containing one cell."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "测试内容"
    return doc


def _get_pPr_children_tags(cell) -> list[str]:
    """Return tag names (without namespace) of all children in cell paragraphs' pPr."""
    tags = []
    for para in cell.paragraphs:
        ppr = para._element.find(qn("w:pPr"))
        if ppr is not None:
            tags = [
                child.tag.split("}")[-1] if "}" in child.tag else child.tag
                for child in ppr
            ]
    return tags


def test_spacing_before_ind_order():
    """_apply_cell_line_spacing 后 w:spacing 应在 w:ind 之前。"""
    doc = _create_table_doc()
    cell = doc.tables[0].cell(0, 0)
    config = _build_config()

    # 先设置对齐（会产生 w:ind）
    _set_cell_alignment(cell, config)
    # 再设行距
    _apply_cell_line_spacing(cell, config)

    tags = _get_pPr_children_tags(cell)
    spacing_idx = tags.index("spacing")
    ind_idx = tags.index("ind")
    assert spacing_idx < ind_idx, (
        f"w:spacing (index {spacing_idx}) must come BEFORE "
        f"w:ind (index {ind_idx}). Actual order: {tags}"
    )


def test_spacing_value_multiple():
    """multiple 模式行距值正确 (1.5 × 240 = 360 twips)。"""
    doc = _create_table_doc()
    cell = doc.tables[0].cell(0, 0)
    config = _build_config(line_spacing=1.5, line_spacing_mode="multiple")

    _apply_cell_line_spacing(cell, config)
    for para in cell.paragraphs:
        ppr = para._element.find(qn("w:pPr"))
        sp = ppr.find(qn("w:spacing"))
        assert sp is not None
        assert sp.get(qn("w:line")) == "360", f"Expected 360, got {sp.get(qn('w:line'))}"
        assert sp.get(qn("w:lineRule")) == "auto", (
            f"Expected auto, got {sp.get(qn('w:lineRule'))}"
        )


def test_spacing_value_fixed():
    """fixed 模式行距值正确 (12pt × 20 = 240 twips)。"""
    doc = _create_table_doc()
    cell = doc.tables[0].cell(0, 0)
    config = _build_config(line_spacing=12.0, line_spacing_mode="fixed", line_spacing_unit="pt")

    _apply_cell_line_spacing(cell, config)
    for para in cell.paragraphs:
        ppr = para._element.find(qn("w:pPr"))
        sp = ppr.find(qn("w:spacing"))
        assert sp is not None
        assert sp.get(qn("w:line")) == "240"
        assert sp.get(qn("w:lineRule")) == "exact"


def test_spacing_overwrite():
    """再次调用 _apply_cell_line_spacing 应覆盖旧值而非追加。"""
    doc = _create_table_doc()
    cell = doc.tables[0].cell(0, 0)

    # 第一次调用：设置 1.5 倍行距
    config1 = _build_config(line_spacing=1.5, line_spacing_mode="multiple")
    _apply_cell_line_spacing(cell, config1)

    # 检查只有 1 个 spacing 元素
    for para in cell.paragraphs:
        ppr = para._element.find(qn("w:pPr"))
        sps = ppr.findall(qn("w:spacing"))
        assert len(sps) == 1, f"Expected 1 spacing element, got {len(sps)}"

    # 第二次调用：改为 2.0 倍
    config2 = _build_config(line_spacing=2.0, line_spacing_mode="multiple")
    _apply_cell_line_spacing(cell, config2)

    for para in cell.paragraphs:
        ppr = para._element.find(qn("w:pPr"))
        sp = ppr.find(qn("w:spacing"))
        assert sp is not None
        assert sp.get(qn("w:line")) == "480", (
            f"Expected 480 for 2.0x, got {sp.get(qn('w:line'))}"
        )
        # 确认仍然只有 1 个 spacing 元素（没有被重复追加）
        sps = ppr.findall(qn("w:spacing"))
        assert len(sps) == 1, (
            f"Expected 1 spacing element after overwrite, got {len(sps)}. "
            f"Old spacing was not cleaned up!"
        )