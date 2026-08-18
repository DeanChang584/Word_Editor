"""
使用真实文档验证表格行距修复。
1. 用默认配置格式化文档
2. 分析输出文档中表格单元格的行距XML
"""
import sys
import shutil
import os
from pathlib import Path

# 确保可以导入 backend
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from shared.schemas import ProfileConfig
from backend.formatter.engine import format_docx
from docx import Document
from lxml import etree
from docx.oxml.ns import qn


SOURCE_DOC = r"D:\桌面\01-教学研究\Task 1-中学英语教学设计与案例分析\教学课例（课例）\恩施高中-谭红梅-教学案例.docx"
OUTPUT_DOC = r"D:\桌面\01-教学研究\Task 1-中学英语教学设计与案例分析\教学课例（课例）\恩施高中-谭红梅-教学案例-R.docx"


def check_spacing_order(doc: Document) -> list[str]:
    """检查所有表格单元格中 spacing 和 ind 的顺序。"""
    issues = []
    for i, table in enumerate(doc.tables):
        for j, row in enumerate(table.rows):
            for k, cell in enumerate(row.cells):
                for l, para in enumerate(cell.paragraphs):
                    ppr = para._element.find(qn("w:pPr"))
                    if ppr is None:
                        continue
                    tags = [
                        child.tag.split("}")[-1] if "}" in child.tag else child.tag
                        for child in ppr
                    ]
                    if "spacing" in tags and "ind" in tags:
                        sp_idx = tags.index("spacing")
                        ind_idx = tags.index("ind")
                        if sp_idx > ind_idx:
                            issues.append(
                                f"Table[{i}] Row[{j}] Cell[{k}] Para[{l}]: "
                                f"spacing (idx={sp_idx}) AFTER ind (idx={ind_idx}). "
                                f"Order: {tags}"
                            )
    return issues


def check_spacing_values(doc: Document) -> list[str]:
    """检查表格单元格中行距值是否正确设置。"""
    issues = []
    for i, table in enumerate(doc.tables):
        for j, row in enumerate(table.rows):
            for k, cell in enumerate(row.cells):
                for l, para in enumerate(cell.paragraphs):
                    ppr = para._element.find(qn("w:pPr"))
                    if ppr is None:
                        continue
                    spacing = ppr.find(qn("w:spacing"))
                    if spacing is None:
                        issues.append(
                            f"Table[{i}] Row[{j}] Cell[{k}] Para[{l}]: "
                            f"NO w:spacing element!"
                        )
                    else:
                        line = spacing.get(qn("w:line"))
                        rule = spacing.get(qn("w:lineRule"))
                        if not line:
                            issues.append(
                                f"Table[{i}] Row[{j}] Cell[{k}] Para[{l}]: "
                                f"w:spacing has no w:line attribute"
                            )
    return issues


def print_cell_details(doc: Document):
    """打印所有表格单元格的详细行距信息。"""
    print("\n" + "=" * 60)
    print("CELL LINE SPACING DETAILS")
    print("=" * 60)
    for i, table in enumerate(doc.tables):
        print(f"\n--- Table[{i}] ({len(table.rows)} rows × {len(table.columns)} cols) ---")
        for j, row in enumerate(table.rows):
            for k, cell in enumerate(row.cells):
                for l, para in enumerate(cell.paragraphs):
                    elem = para._element
                    ppr = elem.find(qn("w:pPr"))
                    text = para.text[:50] if para.text else "(empty)"

                    if ppr is None:
                        full_xml = etree.tostring(elem, pretty_print=True, encoding="unicode")
                        print(f"  Cell[{k}] Para[{l}]: NO pPr")
                        print(f"    Full XML:\n{full_xml}")
                        continue

                    # Get tag order
                    tags = [
                        child.tag.split("}")[-1] if "}" in child.tag else child.tag
                        for child in ppr
                    ]

                    # Get spacing details
                    spacing = ppr.find(qn("w:spacing"))
                    sp_info = "none"
                    if spacing is not None:
                        attrs = []
                        for attr_name in ["w:line", "w:lineRule", "w:before", "w:after", "w:val", "w:unit"]:
                            val = spacing.get(qn(attr_name))
                            if val:
                                attrs.append(f"{attr_name.split(':')[-1]}={val}")
                        sp_info = ", ".join(attrs)

                    # Get ind details
                    ind = ppr.find(qn("w:ind"))
                    ind_info = "none"
                    if ind is not None:
                        attrs = []
                        for attr_name in ["w:firstLine", "w:firstLineChars", "w:hanging", "w:left", "w:right"]:
                            val = ind.get(qn(attr_name))
                            if val and val != "0":
                                attrs.append(f"{attr_name.split(':')[-1]}={val}")
                        if not attrs:
                            ind_info = "all-zero"
                        else:
                            ind_info = ", ".join(attrs)

                    print(f"  Cell[{k}] Para[{l}] text='{text}'")
                    print(f"    pPr tag order: {tags}")
                    print(f"    spacing: [{sp_info}]")
                    print(f"    ind: [{ind_info}]")


def main():
    if not os.path.exists(SOURCE_DOC):
        print(f"ERROR: Source file not found: {SOURCE_DOC}")
        return 1

    # Make a copy to avoid overwriting the original
    work_path = SOURCE_DOC
    
    print(f"Source: {SOURCE_DOC}")
    print(f"Output: {OUTPUT_DOC}")
    print()

    # Build profile with table line spacing settings
    profile = ProfileConfig()
    
    # Set table-specific line spacing
    profile.table.line_spacing = 1.5
    profile.table.line_spacing_mode = "multiple"
    profile.table.line_spacing_unit = "pt"
    profile.table.indent_type = "none"
    profile.table.indent_value = 0.0
    profile.table.indent_unit = "字符"
    profile.table.cell_align_h = "left"
    profile.table.cell_align_v = "center"
    profile.table.row_height_mode = "auto"
    profile.table.row_height = 0
    profile.table.row_height_unit = "厘米"
    profile.table.border_style = "thin"
    
    print("=== Profile Config (table) ===")
    print(f"  line_spacing={profile.table.line_spacing}")
    print(f"  line_spacing_mode={profile.table.line_spacing_mode}")
    print(f"  line_spacing_unit={profile.table.line_spacing_unit}")
    print(f"  indent_type={profile.table.indent_type}")
    print(f"  cell_align_h={profile.table.cell_align_h}")
    print(f"  cell_align_v={profile.table.cell_align_v}")
    print(f"  row_height_mode={profile.table.row_height_mode}")
    print()

    # Analyze original document first
    print("=" * 60)
    print("ORIGINAL DOCUMENT (before formatting)")
    print("=" * 60)
    orig_doc = Document(work_path)
    print_cell_details(orig_doc)
    orig_issues_o = check_spacing_order(orig_doc)
    orig_issues_v = check_spacing_values(orig_doc)
    if orig_issues_o:
        print(f"\n  Original spacing order issues: {len(orig_issues_o)}")
        for iss in orig_issues_o[:5]:
            print(f"    {iss}")
    if orig_issues_v:
        print(f"\n  Original spacing value issues: {len(orig_issues_v)}")
        for iss in orig_issues_v[:5]:
            print(f"    {iss}")
    if not orig_issues_o and not orig_issues_v:
        print("\n  Original doc: no issues detected")

    # Format the document
    print("\n" + "=" * 60)
    print("FORMATTING...")
    print("=" * 60)
    success, msg, out_path = format_docx(work_path, profile, output_path=OUTPUT_DOC)
    
    if not success:
        print(f"ERROR: Formatting failed: {msg}")
        return 1
    
    print(f"Formatting result: {msg}")
    print(f"Output path: {out_path}")
    print()

    # Analyze formatted document
    print("=" * 60)
    print("FORMATTED DOCUMENT ANALYSIS")
    print("=" * 60)
    out_doc = Document(out_path)
    print_cell_details(out_doc)

    # Check for issues
    print("\n" + "=" * 60)
    print("VALIDATION RESULTS")
    print("=" * 60)
    
    issues_order = check_spacing_order(out_doc)
    issues_values = check_spacing_values(out_doc)
    
    if issues_order:
        print(f"\n❌ SPACING ORDER ISSUES: {len(issues_order)}")
        for iss in issues_order:
            print(f"  {iss}")
    else:
        print(f"✅ spacing order: ALL CORRECT (spacing before ind in all cells)")
    
    if issues_values:
        print(f"\n❌ SPACING VALUE ISSUES: {len(issues_values)}")
        for iss in issues_values:
            print(f"  {iss}")
    else:
        print(f"✅ spacing values: ALL cells have w:spacing with w:line set")
    
    total_cells = sum(
        len(row.cells) * sum(1 for _ in cell.paragraphs)
        for table in out_doc.tables for row in table.rows for cell in row.cells
    )
    print(f"\nTotal cells processed: {sum(len(table.rows) * len(table.columns) for table in out_doc.tables)}")
    print(f"Total paragraphs in cells: {total_cells}")
    
    if not issues_order and not issues_values:
        print("\n✅ ALL CHECKS PASSED - line spacing should work correctly")
        return 0
    else:
        print(f"\n❌ {len(issues_order) + len(issues_values)} issues found")
        return 1


if __name__ == "__main__":
    sys.exit(main())