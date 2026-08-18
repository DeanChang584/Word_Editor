"""Analyze table structure and line spacing in a .docx file."""
import sys
import os
from pathlib import Path
from docx import Document
from lxml import etree
from docx.oxml.ns import qn

def analyze_doc(docx_path: str):
    doc = Document(docx_path)
    
    print(f"Document: {docx_path}")
    print(f"Number of tables: {len(doc.tables)}")
    print(f"Number of paragraphs (body): {len(doc.paragraphs)}")
    
    # Check Normal style
    style = doc.styles['Normal']
    pf = style.paragraph_format
    print(f"\nNormal style:")
    print(f"  line_spacing: {pf.line_spacing}")
    print(f"  line_spacing_rule: {pf.line_spacing_rule}")
    print(f"  first_line_indent: {pf.first_line_indent}")
    
    # Analyze tables
    for i, table in enumerate(doc.tables):
        print(f"\n=== Table[{i}] ===")
        print(f"  Rows: {len(table.rows)}, Cols: {len(table.columns)}")
        for j, row in enumerate(table.rows):
            for k, cell in enumerate(row.cells):
                for l, para in enumerate(cell.paragraphs):
                    elem = para._element
                    ppr = elem.find(qn("w:pPr"))
                    
                    # Full XML dump of this paragraph
                    xml_str = etree.tostring(elem, pretty_print=True, encoding="unicode")
                    
                    if ppr is not None:
                        ppr_xml = etree.tostring(ppr, pretty_print=True, encoding="unicode")
                        ind = ppr.find(qn("w:ind"))
                        spacing = ppr.find(qn("w:spacing"))
                        
                        ind_str = "none"
                        if ind is not None:
                            parts = []
                            for attr, val in sorted(ind.attrib.items()):
                                tag = attr.split("}")[-1] if "}" in attr else attr
                                parts.append(f"{tag}={val}")
                            ind_str = ", ".join(parts) if parts else "all-zero"
                        
                        sp_str = "none"
                        if spacing is not None:
                            line = spacing.get(qn("w:line"))
                            before = spacing.get(qn("w:before"))
                            after = spacing.get(qn("w:after"))
                            rule = spacing.get(qn("w:lineRule"))
                            parts = []
                            if line: parts.append(f"line={line}")
                            if before: parts.append(f"before={before}")
                            if after: parts.append(f"after={after}")
                            if rule: parts.append(f"lineRule={rule}")
                            sp_str = ", ".join(parts) if parts else "present(empty)"
                        
                        # Also check for rPr font size
                        rpr = ppr.find(qn("w:rPr"))
                        sz = None
                        if rpr is not None:
                            sz_elem = rpr.find(qn("w:sz"))
                            if sz_elem is not None:
                                sz = sz_elem.get(qn("w:val"))
                        
                        print(f"  Table[{i}] Row[{j}] Cell[{k}] Para[{l}]:")
                        print(f"    indent=[{ind_str}]")
                        print(f"    spacing=[{sp_str}]")
                        print(f"    rPr/szVal={sz}")
                        
                        # Check run-level font size
                        for m, run in enumerate(para.runs):
                            r_elem = run._element
                            rrpr = r_elem.find(qn("w:rPr"))
                            if rrpr is not None:
                                rsz = rrpr.find(qn("w:sz"))
                                rsz_val = rsz.get(qn("w:val")) if rsz is not None else None
                                print(f"    Run[{m}]: szVal={rsz_val}")
                    else:
                        print(f"  Table[{i}] Row[{j}] Cell[{k}] Para[{l}]: NO pPr")
                        print(f"  Full paragraph XML:\n{xml_str}")
    
    # Also check non-table body paragraphs
    print(f"\n=== Body paragraphs (first 5) ===")
    for l, para in enumerate(doc.paragraphs[:5]):
        elem = para._element
        ppr = elem.find(qn("w:pPr"))
        if ppr is not None:
            spacing = ppr.find(qn("w:spacing"))
            sp_str = "none"
            if spacing is not None:
                line = spacing.get(qn("w:line"))
                before = spacing.get(qn("w:before"))
                after = spacing.get(qn("w:after"))
                rule = spacing.get(qn("w:lineRule"))
                parts = []
                if line: parts.append(f"line={line}")
                if before: parts.append(f"before={before}")
                if after: parts.append(f"after={after}")
                if rule: parts.append(f"lineRule={rule}")
                sp_str = ", ".join(parts) if parts else "present(empty)"
            print(f"  Para[{l}] text='{para.text[:40]}' spacing=[{sp_str}]")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        path = sys.argv[1]
    else:
        path = r"D:\桌面\临时Word.docx"
    
    if not os.path.exists(path):
        print(f"ERROR: File not found: {path}")
        sys.exit(1)
    
    analyze_doc(path)